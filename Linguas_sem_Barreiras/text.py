from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

# Criar documento Word para aplicar formatação acadêmica
doc = Document()

# Estilo geral
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p

# Título centralizado
add_heading("Projeto de Intervenção Extensionista", level=1)
add_heading("Raízes do Cerrado no Helizes: Plantio com Nendo Dango (SEMA e IBRAM) para a Preservação", level=2)

# Corpo do relatório
sections = {
    "Data da Realização": "04 de abril de 2025",
    "1. Título do Projeto": "Raízes do Cerrado no Helizes: Plantio com Nendo Dango (SEMA e IBRAM) para a Preservação",
    "2. Área Temática": "Meio Ambiente, Educação Ambiental, Engajamento Comunitário, Ecologia do Cerrado, Técnicas de Plantio Ecológico (Nendo Dango), Parceria Institucional",
    "3. Justificativa": (
        "O presente relatório detalha a execução e os resultados do projeto de intervenção extensionista "
        "\"Raízes do Cerrado no Helizes: Plantio com Nendo Dango (SEMA e IBRAM) para a Preservação\", realizado em "
        "04 de abril de 2025, no Parque Ecológico do Guará (área do Helizes). A iniciativa surgiu do convite da "
        "Secretaria de Estado de Meio Ambiente e Proteção Animal do Distrito Federal (SEMA) e do Instituto Brasília "
        "Ambiental (IBRAM) para colaborar em ações de preservação e engajamento comunitário. A técnica Nendo Dango "
        "foi escolhida como metodologia principal, visando promover uma conexão lúdica e participativa com o processo "
        "de restauração ecológica do Cerrado."
    ),
    "4. Problema Abordado": (
        "A ação buscou aumentar o engajamento da comunidade do Guará nas iniciativas de preservação do Parque Ecológico "
        "do Guará (área do Helizes), em colaboração com a SEMA e o IBRAM, utilizando o método Nendo Dango para o plantio "
        "de espécies nativas do Cerrado, promovendo a educação ambiental sobre o bioma e incentivando a participação ativa "
        "na restauração ecológica, em consonância com as diretrizes dos órgãos ambientais."
    )
}

# Adicionar seções fixas
for title, content in sections.items():
    add_paragraph(f"{title}:", bold=True)
    add_paragraph(content)

# Salvar como DOCX temporário
docx_path = "/mnt/data/Projeto_Raizes_do_Cerrado.docx"
pdf_path = "/mnt/data/Projeto_Raizes_do_Cerrado.pdf"
doc.save(docx_path)

# Gerar PDF básico com FPDF
pdf = FPDF()
pdf.add_page()
pdf.set_font("Times", size=12)

with open(docx_path, "r", encoding="utf-8", errors="ignore") as f:
    for line in f:
        pdf.multi_cell(0, 10, txt=line.strip())

pdf.output(pdf_path)
pdf_path
